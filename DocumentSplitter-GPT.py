import os
from docx import Document

def split_document(input_path, output_path_template, words_per_file=18000):
    # Create the output directory if it doesn't exist
    output_dir = os.path.dirname(output_path_template)
    os.makedirs(output_dir, exist_ok=True)

    # Load the input document
    doc = Document(input_path)

    # Initialize variables
    current_doc = Document()
    word_count = 0
    file_index = 1

    # Function to save the current document and start a new one
    def save_current_document():
        nonlocal current_doc, file_index
        output_path = output_path_template.format(file_index)
        current_doc.save(output_path)
        print(f'Saved: {output_path}')
        current_doc = Document()
        file_index += 1

    # Iterate through paragraphs and add them to the current document
    for paragraph in doc.paragraphs:
        words_in_paragraph = len(paragraph.text.split())
        if word_count + words_in_paragraph > words_per_file:
            # Save the current document and start a new one
            save_current_document()
            word_count = 0
        
        # Add the paragraph to the current document
        current_doc.add_paragraph(paragraph.text)
        word_count += words_in_paragraph

    # Save any remaining paragraphs in the last document
    if word_count > 0:
        save_current_document()

# Example usage
input_file_path = 'path/to/your/input/document.docx'
output_file_path_template = 'path/to/your/output/directory/document_part_{}.docx'
split_document(input_file_path, output_file_path_template)
